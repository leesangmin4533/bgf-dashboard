#!/usr/bin/env python3
"""
사업계획서 docx 유효성 검사 스크립트
- 구조, 필수 키워드, 표 수, 스타일 검증
"""
import sys
import os
from docx import Document


def validate(path):
    errors = []
    warnings = []

    if not os.path.exists(path):
        print(f"[FAIL] 파일 없음: {path}")
        return False

    doc = Document(path)

    # 1. 기본 구조 검증
    paragraphs = doc.paragraphs
    tables = doc.tables

    total_paras = len(paragraphs)
    total_tables = len(tables)

    print(f"  단락 수: {total_paras}")
    print(f"  표 수: {total_tables}")

    if total_paras < 50:
        errors.append(f"단락 수 부족: {total_paras}개 (최소 50개 기대)")
    if total_tables < 15:
        errors.append(f"표 수 부족: {total_tables}개 (최소 15개 기대)")

    # 2. 필수 키워드 검증
    full_text = "\n".join([p.text for p in paragraphs])
    # 표 안의 텍스트도 수집
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text

    required_keywords = [
        ("25개 Feature", "ML Feature 수 (22 → 25 수정 여부)"),
        ("카카오톡 나에게 보내기 API", "카카오 알림 설명"),
        ("9단계", "9단계 예측 파이프라인"),
        ("WMA", "가중이동평균"),
        ("Feature 블렌딩", "Feature 블렌딩 단계"),
        ("ML 앙상블", "ML 앙상블 보정"),
        ("FoodStrategy", "카테고리 Strategy 클래스명"),
        ("BeerStrategy", "맥주 Strategy"),
        ("TobaccoStrategy", "담배 Strategy"),
        ("금요일 2.54배", "맥주 요일 패턴 수치"),
        ("common.db", "공통 DB 파일명"),
        ("stores/", "매장별 DB 경로"),
        ("DBRouter", "DB 라우터 클래스명"),
        ("PreOrderEvaluator", "사전 평가 모듈명"),
        ("FORCE_ORDER", "강제발주 판정"),
        ("일평균×0.59", "인기도 가중치"),
        ("ThreadPoolExecutor", "병렬 실행 기술"),
        ("MultiStoreRunner", "병렬 실행 클래스명"),
        ("StoreContext", "매장 컨텍스트"),
        ("frozen dataclass", "스레드 안전 패턴"),
        ("넥사크로", "넥사크로 플랫폼"),
        ("JavaScript", "넥사크로 접근 방식"),
        ("ActionChains", "Selenium 입력 방식"),
        ("예보 기온", "날씨 예보 연동"),
        ("1,268개 테스트", "테스트 수"),
        ("new_product", "신상품 모듈"),
        ("홈 / 예측분석 / 발주 / 리포트 / 규칙", "웹 대시보드 5개 탭"),
        ("발주완료", "카카오 알림 종류 1"),
        ("폐기위험", "카카오 알림 종류 2"),
        ("주간리포트", "카카오 알림 종류 5"),
    ]

    pass_count = 0
    for keyword, desc in required_keywords:
        if keyword in full_text:
            pass_count += 1
        else:
            errors.append(f"필수 키워드 누락: '{keyword}' ({desc})")

    print(f"  필수 키워드: {pass_count}/{len(required_keywords)} 통과")

    # 3. 문서 구조 (헤딩) 검증
    headings = [p.text for p in paragraphs if p.style.name.startswith("Heading")]
    required_sections = [
        "일반현황",
        "창업아이템 개요",
        "문제인식",
        "실현가능성",
        "성장전략",
        "조직구성",
    ]

    for section in required_sections:
        found = any(section in h for h in headings)
        if not found:
            errors.append(f"필수 섹션 누락: '{section}'")

    print(f"  필수 섹션: {len(required_sections) - sum(1 for s in required_sections if any(s in h for h in headings))} 누락")

    # 4. 표지 검증
    if "사 업 계 획 서" not in full_text:
        errors.append("표지 '사 업 계 획 서' 타이틀 누락")

    if "창업중심대학 지역기반" not in full_text:
        errors.append("표지 '창업중심대학 지역기반' 누락")

    # 5. [대표자명] 등 플레이스홀더 유지 확인
    placeholders = ["[대표자명]", "[법인명]"]
    for ph in placeholders:
        if ph not in full_text:
            warnings.append(f"플레이스홀더 '{ph}'가 없음 (미확정 정보는 괄호 표기 유지 필요)")

    # 6. 오류 수정 검증 (틀린 값이 남아있지 않은지)
    wrong_values = [
        ("22개 Feature", "25개 Feature로 수정 필요"),
        ("카카오톡 5종 알림", "카카오톡 나에게 보내기 API로 수정 필요"),
        ("웹 대시보드 5개 페이지", "웹 대시보드 5개 탭으로 수정 필요"),
        ("아키텍처", "규칙 현황판으로 대체 필요 (탭명)"),
    ]

    for wrong, fix in wrong_values:
        if wrong in full_text:
            # "아키텍처" 는 시스템 아키텍처 설명에서 사용 가능 → 탭 이름으로서만 확인
            if wrong == "아키텍처":
                # 5개 페이지: ... 아키텍처 패턴만 확인
                if "홈, 발주관리, 예측분석, 리포트, 아키텍처" in full_text:
                    errors.append(f"오류 미수정: '{wrong}' ({fix})")
            else:
                errors.append(f"오류 미수정: '{wrong}' ({fix})")

    # ── 결과 출력 ──
    print()
    if errors:
        print(f"[FAIL] 검증 실패 — {len(errors)}개 오류:")
        for e in errors:
            print(f"  ✗ {e}")
    else:
        print("[PASS] 모든 검증 통과!")

    if warnings:
        print(f"\n[WARN] {len(warnings)}개 경고:")
        for w in warnings:
            print(f"  ⚠ {w}")

    return len(errors) == 0


def main():
    if len(sys.argv) > 1:
        path = sys.argv[1]
    else:
        # 기본 경로 탐색
        candidates = [
            "/mnt/user-data/outputs/사업계획서_창업중심대학_지역기반_v4.docx",
            os.path.join(os.path.dirname(__file__), "..", "..",
                         "사업계획서_창업중심대학_지역기반_v4.docx"),
        ]
        path = None
        for c in candidates:
            c = os.path.normpath(c)
            if os.path.exists(c):
                path = c
                break

        if not path:
            print("[ERROR] 사업계획서 파일을 찾을 수 없습니다.")
            print("  사용법: python validate.py <파일경로>")
            sys.exit(1)

    print(f"검증 대상: {path}")
    print("=" * 60)

    success = validate(path)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
