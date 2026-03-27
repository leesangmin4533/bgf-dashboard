#!/usr/bin/env python3
"""
사업계획서 (창업중심대학 지역기반) v4 생성 스크립트
- 실제 시스템 스펙 기반으로 재작성
- python-docx 라이브러리 사용
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ────────────────────────────────────────────
# 설정
# ────────────────────────────────────────────
FONT_NAME = "맑은 고딕"
HEADER_BG = "D6E4F0"
MARGIN_DXA = 1000  # ~1.76cm


def set_cell_bg(cell, color):
    """셀 배경색 설정"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_font(run, size=10, bold=False, color=None, name=FONT_NAME):
    """폰트 설정"""
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    """스타일 적용된 헤딩"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    return h


def add_para(doc, text, size=10, bold=False, align=None, space_after=6):
    """일반 단락 추가"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """표 생성 (헤더 배경색 적용)"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 행
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=9, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, HEADER_BG)

    # 데이터 행
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=9)

    # 열 너비 설정
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 표 뒤 간격
    return table


def add_bullet(doc, text, level=0, size=10):
    """불릿 단락"""
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    set_font(run, size=size)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p


# ════════════════════════════════════════════
# 메인 생성 함수
# ════════════════════════════════════════════
def create_document():
    doc = Document()

    # ── 페이지 설정 ──
    for section in doc.sections:
        section.top_margin = Emu(MARGIN_DXA * 914)
        section.bottom_margin = Emu(MARGIN_DXA * 914)
        section.left_margin = Emu(MARGIN_DXA * 914)
        section.right_margin = Emu(MARGIN_DXA * 914)

    # ════════════════════════════════════════
    # 표지
    # ════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("사 업 계 획 서")
    set_font(run, size=28, bold=True)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AI 기반 편의점 자동 발주 시스템")
    set_font(run, size=16, bold=True, color=(0, 51, 102))

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run("(창업중심대학 지역기반)")
    set_font(run, size=14, color=(80, 80, 80))

    for _ in range(4):
        doc.add_paragraph()

    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("사 업 명", "AI 기반 편의점 자동 발주 시스템"),
        ("대 표 자", "[대표자명]"),
        ("법 인 명", "[법인명]"),
        ("사업기간", "2026. 03 ~ 2026. 12 (10개월)"),
    ]
    for i, (label, value) in enumerate(info_data):
        cell_l = info_table.rows[i].cells[0]
        cell_l.text = ""
        run = cell_l.paragraphs[0].add_run(label)
        set_font(run, size=11, bold=True)
        cell_l.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell_l, HEADER_BG)
        cell_l.width = Cm(4)

        cell_r = info_table.rows[i].cells[1]
        cell_r.text = ""
        run = cell_r.paragraphs[0].add_run(value)
        set_font(run, size=11)
        cell_r.width = Cm(10)

    doc.add_page_break()

    # ════════════════════════════════════════
    # □ 일반현황
    # ════════════════════════════════════════
    add_heading_styled(doc, "일반현황", level=1)

    add_table(doc,
        ["항목", "내용"],
        [
            ["대표자", "[대표자명]"],
            ["법인명", "[법인명]"],
            ["사업자등록번호", "[사업자등록번호]"],
            ["주소", "경기도 이천시"],
            ["연락처", "[연락처]"],
            ["이메일", "[이메일]"],
            ["업종", "소프트웨어 개발 / AI 솔루션"],
            ["사업 형태", "예비 창업자 (개인 개발)"],
            ["창업 예정일", "2026년 상반기"],
        ],
        col_widths=[4, 12]
    )

    doc.add_page_break()

    # ════════════════════════════════════════
    # □ 창업아이템 개요 (요약)
    # ════════════════════════════════════════
    add_heading_styled(doc, "창업아이템 개요 (요약)", level=1)

    add_para(doc,
        "본 사업은 CU 편의점 운영 현장에서 직접 개발한 "
        "'AI 기반 자동 발주 시스템'을 상용화하는 것을 목표로 합니다. "
        "현재 3개 매장(이천호반베르디움점, 이천동양점, 마평로드점)에서 "
        "실전 운영 중이며, 15개 카테고리별 맞춤 예측 + ML 앙상블 + "
        "자동 발주 실행까지 완전 자동화된 시스템입니다.",
        size=10)

    add_table(doc,
        ["구분", "내용"],
        [
            ["제품명", "AI 기반 편의점 자동 발주 시스템 (BGF Auto Order)"],
            ["핵심 기술",
             "9단계 예측 파이프라인 (WMA + Feature 블렌딩 + 6종 계수 + ML 앙상블)\n"
             "15개 카테고리별 독립 Strategy 패턴\n"
             "넥사크로(Nexacro) 기반 자동 발주 실행"],
            ["기술 스택",
             "Python 3.12 / Selenium / SQLite / Flask / scikit-learn\n"
             "RF + GradientBoosting 앙상블 (25개 Feature Engineering)"],
            ["적용 매장", "3개 매장 실전 운영 (이천호반베르디움점, 이천동양점, 마평로드점)"],
            ["테스트 검증", "1,268개 테스트 통과 (자동화 테스트 스위트)"],
            ["알림 시스템",
             "카카오톡 나에게 보내기 API (OAuth 기반) 5종 알림\n"
             "발주완료 / 폐기위험 / 일일리포트 / 행사변경 / 주간리포트"],
            ["웹 대시보드",
             "Flask 기반 5개 탭: 홈 / 예측분석 / 발주 / 리포트 / 규칙 현황판\n"
             "96개 API 엔드포인트, 실시간 모니터링"],
            ["차별점",
             "실 매장 운영 데이터 기반 검증 완료\n"
             "넥사크로 플랫폼 직접 제어 (업계 유일)\n"
             "카테고리별 독립 전략으로 도시락~담배까지 전품목 커버"],
        ],
        col_widths=[3.5, 12.5]
    )

    doc.add_page_break()

    # ════════════════════════════════════════
    # 1. 문제인식
    # ════════════════════════════════════════
    add_heading_styled(doc, "1. 문제인식", level=1)

    # 1-1. 창업 배경 및 동기
    add_heading_styled(doc, "1-1. 창업 배경 및 동기", level=2)

    add_para(doc,
        "편의점 발주는 점주의 경험과 감에 의존하는 대표적인 수작업 영역입니다. "
        "CU(BGF 리테일) 기준 전국 17,000여 매장의 점주들이 매일 300~1,500개 "
        "상품의 발주량을 수동으로 결정하고 있으며, 이로 인해 다음과 같은 구조적 "
        "문제가 발생합니다.", size=10)

    add_table(doc,
        ["문제 유형", "현황", "경제적 영향"],
        [
            ["폐기 손실",
             "유통기한 1~3일 신선식품(도시락, 김밥 등)\n과잉 발주 시 전량 폐기",
             "매장당 월 30~80만원 손실\n(연 360~960만원)"],
            ["품절 기회손실",
             "과소 발주 시 매대 공백 발생\n특히 맥주(금요일), 아이스크림(여름) 등",
             "매출 기회 10~15% 상실\n고객 이탈 가속"],
            ["점주 시간 낭비",
             "매일 1~2시간 수동 발주 작업\n숙련 점주도 경험적 판단에 의존",
             "인건비 환산 월 50~100만원\n다매장 운영 시 병목"],
            ["데이터 미활용",
             "판매 데이터 축적되지만 분석 없이 방치\n요일/계절/행사 패턴 미반영",
             "최적화 불가능\n반복적 실수"],
        ],
        col_widths=[3, 6.5, 6.5]
    )

    add_para(doc,
        "본 대표자는 CU 편의점을 직접 운영하면서 이 문제를 체감하였고, "
        "2026년 1월부터 Python 기반의 자동 발주 시스템을 독자 개발하여 "
        "현재 3개 매장에서 실전 운영 중입니다. "
        "본 사업계획서는 이 검증된 시스템을 상용화하기 위한 계획입니다.",
        size=10)

    # 1-2. 창업 아이템의 혁신성 및 차별성
    add_heading_styled(doc, "1-2. 창업 아이템의 혁신성 및 차별성", level=2)

    add_para(doc,
        "기존 BGF 리테일 본사 제공 '스마트발주'는 과거 판매량 단순 평균 기반으로 "
        "카테고리별 특성(도시락 유통기한 1일 vs 담배 보루 단위 등)을 반영하지 못합니다. "
        "본 시스템은 다음과 같은 기술적 혁신을 통해 차별화됩니다.",
        size=10)

    add_table(doc,
        ["차별화 요소", "기존 (BGF 스마트발주)", "본 시스템 (AI Auto Order)"],
        [
            ["예측 모델",
             "과거 N일 단순 평균",
             "9단계 예측 파이프라인\nWMA + Feature 블렌딩 + ML 앙상블"],
            ["카테고리 대응",
             "전품목 동일 로직",
             "15개 카테고리별 독립 Strategy\n(도시락/맥주/담배 등 각각 다른 알고리즘)"],
            ["외부 요인 반영",
             "미반영",
             "기온 예보 / 요일 / 공휴일 / 계절 / 트렌드\n6종 계수 자동 적용"],
            ["발주 실행",
             "수동 확인 후 수정",
             "넥사크로 직접 제어로 완전 자동 입력"],
            ["다매장 지원",
             "매장별 개별 접속",
             "ThreadPoolExecutor 병렬 실행\n3개 매장 동시 처리"],
            ["피드백 루프",
             "없음",
             "예측 vs 실적 비교 → 자동 보정\nML 매일 증분학습"],
        ],
        col_widths=[3, 5.5, 7.5]
    )

    doc.add_page_break()

    # ════════════════════════════════════════
    # 2. 실현가능성
    # ════════════════════════════════════════
    add_heading_styled(doc, "2. 실현가능성", level=1)

    # 2-1. 기술적 실현가능성
    add_heading_styled(doc, "2-1. 기술적 실현가능성 (구현 완료 시스템 상세)", level=2)

    add_para(doc, "■ 시스템 아키텍처", size=11, bold=True, space_after=4)
    add_para(doc,
        "5계층 클린 아키텍처 기반으로 설계되었으며, 각 계층은 독립적으로 "
        "테스트 가능합니다.",
        size=10)

    add_table(doc,
        ["계층", "패키지", "역할", "주요 모듈"],
        [
            ["Settings", "src/settings/",
             "전역 설정, 상수 관리",
             "app_config, constants, timing\nStoreContext (frozen dataclass)"],
            ["Domain", "src/domain/",
             "순수 비즈니스 로직 (I/O 없음)",
             "CategoryStrategy (15개)\nPredictionResult, OrderItem\nPreOrderEvaluator"],
            ["Infrastructure", "src/infrastructure/",
             "모든 I/O 처리",
             "DBRouter (자동 라우팅)\n19개 Repository\nBGF 넥사크로 드라이버\nKakaoNotifier"],
            ["Application", "src/application/",
             "오케스트레이션",
             "DailyOrderFlow (메인)\nMultiStoreRunner (병렬)\n7개 Use Case"],
            ["Presentation", "src/presentation/",
             "사용자 인터페이스",
             "Flask 웹 대시보드 (96개 API)\nCLI (argparse subcommands)"],
        ],
        col_widths=[2.5, 3, 4, 6.5]
    )

    # ── 9단계 예측 파이프라인 ──
    add_para(doc, "■ 9단계 예측 파이프라인 (ImprovedPredictor)", size=11, bold=True, space_after=4)
    add_para(doc,
        "핵심 예측 엔진은 ImprovedPredictor 클래스로, 9단계 순차 파이프라인으로 "
        "최종 발주량을 산출합니다.",
        size=10)

    add_table(doc,
        ["단계", "모듈", "설명", "상세"],
        [
            ["1단계", "WMA\n(가중이동평균)",
             "최근 판매 데이터 기반 기본 예측",
             "어제 25% / 2일전 20% / 3일전 15%\n/ 4~7일전 각 10% (합계 40%)"],
            ["2단계", "Feature 블렌딩",
             "통계적 보정",
             "EWM (지수가중이동평균)\n+ 동요일 평균 블렌딩"],
            ["3단계", "휴일 계수",
             "공휴일/연휴 수요 반영",
             "공공데이터 API 연동\n연휴 전날/당일/다음날 차등"],
            ["4단계", "기온 계수",
             "배송일 기온 기반 조정",
             "예보 기온 우선 > 실측 > 1.0 폴백\n음료 30도↑ +15%, 도시락 30도↑ -10%"],
            ["5단계", "요일 계수",
             "카테고리별 요일 패턴",
             "예: 맥주 금요일 2.54배\n과거 4주 요일별 평균 자동 산출"],
            ["6단계", "계절 계수",
             "7그룹 월별 계절성",
             "아이스크림 여름 1.50배/겨울 0.60배\n음료/라면/주류 등 그룹별 차등"],
            ["7단계", "트렌드 조정",
             "상승/하락 추세 반영",
             "2주 vs 4주 평균 비교\n조정 범위: ±8~15%"],
            ["8단계", "안전재고",
             "카테고리별 Strategy 적용",
             "15개 Strategy 클래스\n유통기한/회전율/계절성 반영"],
            ["9단계", "ML 앙상블 보정",
             "머신러닝 최종 보정",
             "RF + GradientBoosting\n25개 Feature Engineering\nMAE 기반 적응형 블렌딩(0.1~0.5)"],
        ],
        col_widths=[1.5, 2.5, 3.5, 8.5]
    )

    # ── 15개 카테고리 전략 ──
    add_para(doc, "■ 15개 카테고리별 독립 전략 (Strategy 패턴)", size=11, bold=True, space_after=4)
    add_para(doc,
        "각 상품 카테고리는 고유한 수요 특성을 가지므로 독립된 Strategy 클래스로 "
        "분리하여 관리합니다. Strategy 패턴(GoF)을 적용하여 새 카테고리 추가 시 "
        "기존 코드 수정 없이 확장 가능합니다.",
        size=10)

    add_table(doc,
        ["Strategy 클래스", "대상 중분류(mid_cd)", "핵심 로직"],
        [
            ["FoodStrategy",
             "001(도시락), 002(주먹밥)\n003(김밥), 004(샌드위치)\n005(햄버거), 012(빵)",
             "유통기한 1~3일 동적 폐기계수\n일별 판매 기반 안전재고\n시간대별 배송 수요 비율 적용"],
            ["BeerStrategy",
             "049(맥주)",
             "요일 패턴 — 금요일 2.54배\n주말 전 3일 안전재고\n기온 30도↑ 추가 15%"],
            ["SojuStrategy",
             "050(소주)",
             "요일 패턴 + 겨울 상승 계수\n맥주와 연동 분석"],
            ["TobaccoStrategy",
             "072(담배), 073(전자담배)",
             "보루 단위(10갑) + 전량소진 패턴\nmax_order=30, 안정 수요"],
            ["BeverageStrategy",
             "040~048(음료류)",
             "온도/계절 영향 (여름 +30%)\n기온 예보 연동"],
            ["FrozenIceStrategy",
             "034(아이스크림)\n035(냉동식품), 100",
             "계절성 극대 (여름 1.50배/겨울 0.60배)\n냉동고 용량 제한 반영"],
            ["RamenStrategy",
             "006(조리면), 032(면류)",
             "회전율 기반 안전재고\n겨울 수요 상승"],
            ["SnackConfectionStrategy",
             "014~020(과자/제과)\n029~030(간식류)",
             "프로모션 민감도 높음\n행사 배수 보정"],
            ["PerishableStrategy",
             "013(유제품), 026(반찬)\n046(신선식품)",
             "요일 가중 + 유통기한 관리"],
            ["InstantMealStrategy",
             "027~028(즉석밥/국)\n031,033,035(즉석조리)",
             "비상 수요 패턴\n재난/한파 시 급등 대응"],
            ["AlcoholGeneralStrategy",
             "052(양주), 053(와인)",
             "저회전 상품 관리\n최소 재고 유지"],
            ["DailyNecessityStrategy",
             "036~037(세면/위생)\n056~057(주방/생활)\n086(세탁용품)",
             "안정 수요, 장기 안전재고"],
            ["GeneralMerchandiseStrategy",
             "054~071(잡화/비식품)",
             "안정 수요, 기본 예측"],
            ["DessertStrategy",
             "014(디저트)",
             "발주 유지/정지 자동 판단\n4카테고리 분류, 3생애주기"],
            ["DefaultStrategy",
             "900(소모품) 및 미분류",
             "기본 WMA 예측\n범용 안전재고"],
        ],
        col_widths=[3.5, 4.5, 8]
    )

    # ── DB 구조 ──
    add_para(doc, "■ 데이터베이스 이중 분할 구조", size=11, bold=True, space_after=4)
    add_para(doc,
        "공통 DB(common.db)와 매장별 DB(stores/{store_id}.db)의 이중 구조로 "
        "설계되었습니다. DBRouter가 테이블 종류에 따라 자동으로 올바른 DB에 "
        "라우팅합니다.",
        size=10)

    add_table(doc,
        ["DB 구분", "파일", "주요 테이블", "역할"],
        [
            ["공통 DB", "data/common.db",
             "products (상품 마스터)\nmid_categories (중분류)\nproduct_details (상세/마진)\nexternal_factors (날씨/공휴일)\nstores, app_settings 등 8개",
             "매장 무관 전역 참조 데이터\n상품 정보, 카테고리 체계\n날씨/공휴일 외부 요인"],
            ["매장별 DB", "data/stores/46513.db\ndata/stores/46704.db\ndata/stores/47863.db",
             "daily_sales (일별 판매)\norder_tracking (발주 추적)\ninventory_batches (FIFO 배치)\nrealtime_inventory (실시간 재고)\nprediction_logs, eval_outcomes\npromotions 등 15개",
             "매장 고유 운영 데이터\n판매/재고/발주/예측 이력\n매장간 완전 격리"],
            ["DB 라우터", "DBRouter 클래스",
             "자동 라우팅 로직\nATTACH DATABASE 교차참조",
             "테이블명 기반 자동 DB 선택\n매장 DB에서 공통 DB 참조 시\nATTACH 패턴 사용"],
        ],
        col_widths=[2.5, 4, 5, 4.5]
    )

    # ── 사전 발주 평가 ──
    add_para(doc, "■ 사전 발주 평가 (PreOrderEvaluator)", size=11, bold=True, space_after=4)
    add_para(doc,
        "모든 상품은 발주 전에 5단계 판정을 거쳐 발주 긴급도를 결정합니다. "
        "인기도 점수는 다차원 가중 합산으로 산출됩니다.",
        size=10)

    add_table(doc,
        ["판정", "조건", "처리"],
        [
            ["FORCE_ORDER\n(강제발주)",
             "현재고 + 미입고 = 0\n(단, 일평균 < 0.1이면 NORMAL 다운그레이드)",
             "즉시 발주 실행\n품절 방지 최우선"],
            ["URGENT\n(긴급)",
             "노출일 < 1.0일 AND 중인기 이상\n(인기도 = 일평균×0.59 + 판매일비율×0.27\n + 트렌드×0.14)",
             "우선 발주\n안전재고 상향"],
            ["NORMAL\n(일반)",
             "노출일 < 2.0일",
             "정상 발주\n표준 안전재고"],
            ["SKIP\n(생략)",
             "노출일 > 2.5일 AND 저인기",
             "발주 생략\n재고 충분"],
            ["PASS\n(보류)",
             "위 조건에 해당하지 않는 경우",
             "발주 보류\n다음 회차 재평가"],
        ],
        col_widths=[3, 7, 6]
    )

    # ── 병렬 실행 ──
    add_para(doc, "■ 다매장 병렬 실행", size=11, bold=True, space_after=4)
    add_para(doc,
        "MultiStoreRunner 클래스가 ThreadPoolExecutor(max_workers=4) 기반으로 "
        "다매장 작업을 병렬 처리합니다. StoreContext(frozen dataclass)로 "
        "멀티스레드 안전성을 보장하며, 매장별 DB가 물리적으로 분리되어 있어 "
        "잠금(Lock) 경합이 발생하지 않습니다.",
        size=10)

    # ── 날씨 예보 연동 ──
    add_para(doc, "■ 날씨 예보 연동", size=11, bold=True, space_after=4)
    add_para(doc,
        "배송일(내일) 예보 기온을 BGF 시스템에서 자동 수집하여 발주량을 조정합니다. "
        "3단계 폴백 구조: 예보 기온(우선) → 실측 기온 → 1.0(기본값). "
        "카테고리별 차등 적용: 음료 30도 이상 +15%, 도시락 30도 이상 -10% 등.",
        size=10)

    # ── 발주 실행 방식 ──
    add_para(doc, "■ 발주 실행 방식 (넥사크로 직접 제어)", size=11, bold=True, space_after=4)
    add_para(doc,
        "BGF 스토어 시스템은 넥사크로(Nexacro) 플랫폼 기반으로 일반 CSS/XPath "
        "선택자가 작동하지 않습니다. JavaScript로 넥사크로 내부 객체에 직접 접근하여 "
        "데이터셋을 조작하고, Selenium ActionChains(Ctrl+A → Delete → send_keys)로 "
        "값을 입력합니다. 3단계 폴백 구조: Direct API → Batch Grid → Selenium.",
        size=10)

    # ── 신상품 도입 현황 ──
    add_para(doc, "■ 신상품 도입 현황 자동 관리", size=11, bold=True, space_after=4)
    add_para(doc,
        "BGF 상생지원제도의 신상품 도입 지원금 극대화를 위해 월별 도입률과 "
        "3일발주 달성률을 자동 추적합니다. 종합점수 95점 이상 달성 시 "
        "월 160,000원 지원금 수령 가능. DB 3개 테이블(new_product_status, "
        "new_product_items, new_product_monthly)로 관리.",
        size=10)

    doc.add_page_break()

    # 2-2. 시장 실현가능성
    add_heading_styled(doc, "2-2. 시장 실현가능성", level=2)

    add_para(doc,
        "편의점 자동 발주 시장은 아직 초기 단계이며, 현재 시장에 "
        "점주 관점의 AI 발주 솔루션은 사실상 부재합니다.",
        size=10)

    add_table(doc,
        ["시장 지표", "수치", "출처"],
        [
            ["국내 편의점 수", "약 55,000개 (2025)", "산업통상자원부"],
            ["CU 매장 수", "약 17,000개", "BGF 리테일 IR"],
            ["매장당 발주 상품 수", "300~1,500개/일", "현장 운영 경험"],
            ["폐기 손실 (신선식품)", "매장당 월 30~80만원", "현장 운영 데이터"],
            ["점주 발주 소요 시간", "매일 1~2시간", "현장 설문"],
            ["타겟 시장 규모", "CU 17,000개 × 월 3~5만원\n= 연 61~102억원", "추정"],
        ],
        col_widths=[4, 5, 7]
    )

    # ── 최종 산출물 표 ──
    add_para(doc, "■ 최종 산출물 현황표", size=11, bold=True, space_after=4)

    add_table(doc,
        ["산출물", "현황", "세부 내역"],
        [
            ["예측 엔진", "구현 완료\n(운영 중)",
             "ImprovedPredictor 9단계 파이프라인\n15개 카테고리 Strategy\n"
             "RF+GB ML 앙상블 (25개 Feature)\n수요패턴 4단계 분류 (DemandClassifier)"],
            ["자동 발주", "구현 완료\n(운영 중)",
             "넥사크로 JavaScript 직접 제어\n3단계 폴백: Direct API → Batch Grid → Selenium\n"
             "PreOrderEvaluator 5단계 판정\n행사(1+1, 2+1) 배수 자동 보정"],
            ["웹 대시보드", "구현 완료\n(운영 중)",
             "Flask 기반 5개 탭 (홈/예측분석/발주/리포트/규칙)\n"
             "96개 API 엔드포인트\n실시간 재고/예측/발주 모니터링"],
            ["알림 시스템", "구현 완료\n(운영 중)",
             "카카오톡 나에게 보내기 API (OAuth)\n"
             "5종: 발주완료/폐기위험/일일리포트/행사변경/주간리포트"],
            ["DB 시스템", "구현 완료\n(운영 중)",
             "SQLite 이중 분할 구조\n공통 DB (8개 테이블) + 매장별 DB (15개 테이블)\n"
             "DBRouter 자동 라우팅, ATTACH 교차참조"],
            ["병렬 실행", "구현 완료\n(운영 중)",
             "MultiStoreRunner (ThreadPoolExecutor, max_workers=4)\n"
             "StoreContext (frozen dataclass) 스레드 안전\n3개 매장 동시 처리"],
            ["테스트 스위트", "구현 완료",
             "1,268개 테스트 통과\npytest 자동화\n회귀 방지 + CI 연동 가능"],
            ["날씨 연동", "구현 완료\n(운영 중)",
             "BGF 예보 기온 자동 수집\n3단계 폴백: 예보 > 실측 > 기본값\n"
             "카테고리별 기온 교차 계수"],
            ["신상품 관리", "구현 완료\n(운영 중)",
             "도입률/달성률 자동 추적\n3일발주 스케줄링\n"
             "월 160,000원 지원금 최적화"],
        ],
        col_widths=[3, 2.5, 10.5]
    )

    doc.add_page_break()

    # ════════════════════════════════════════
    # 3. 성장전략
    # ════════════════════════════════════════
    add_heading_styled(doc, "3. 성장전략", level=1)

    # 3-1. 사업화 전략
    add_heading_styled(doc, "3-1. 사업화 전략", level=2)

    add_para(doc,
        "현재 자체 매장에서 검증 완료된 시스템을 3단계로 확장합니다.",
        size=10)

    add_table(doc,
        ["단계", "기간", "목표", "핵심 과제"],
        [
            ["1단계\nPoC 검증",
             "2026.03 ~ 06\n(4개월)",
             "자체 3개 매장 안정 운영\n성과 데이터 축적",
             "폐기율 30% 감소 실증\n품절률 50% 감소 실증\n시스템 안정성 99% 확보"],
            ["2단계\nBeta 확장",
             "2026.07 ~ 09\n(3개월)",
             "외부 10개 매장 확장\nSaaS 모델 검증",
             "매장 온보딩 자동화\n클라우드 인프라 구축\n월정액 과금 모델 테스트"],
            ["3단계\n상용화",
             "2026.10 ~ 12\n(3개월)",
             "50개 매장 서비스\n매출 목표 달성",
             "영업 채널 구축\n고객 지원 체계\nGS25/세븐일레븐 확장 준비"],
        ],
        col_widths=[2.5, 3, 4.5, 6]
    )

    # 3-2. 수익 모델
    add_heading_styled(doc, "3-2. 수익 모델", level=2)

    add_table(doc,
        ["수익원", "가격", "설명"],
        [
            ["Basic 플랜\n(소규모 매장)", "월 30,000원/매장",
             "자동 발주 + 기본 예측\n카카오톡 알림\n웹 대시보드"],
            ["Pro 플랜\n(다매장 운영)", "월 50,000원/매장",
             "Basic 전체 + ML 앙상블\n15개 카테고리 전략\n신상품 관리 + 폐기 분석"],
            ["Enterprise\n(본사/경영주)", "별도 협의",
             "멀티브랜드 대응\n맞춤 분석 리포트\nAPI 연동"],
        ],
        col_widths=[3, 3.5, 9.5]
    )

    # 3-3. 자금 운용 계획
    add_heading_styled(doc, "3-3. 자금 운용 계획", level=2)

    # 3-3-1
    add_heading_styled(doc, "3-3-1. 정부지원금 사용 계획", level=3)

    add_table(doc,
        ["항목", "금액(만원)", "용도", "비고"],
        [
            ["클라우드 인프라", "800",
             "AWS/NCP 서버 (멀티테넌트)\nDB 서버, 백업 스토리지",
             "월 80만원 × 10개월"],
            ["개발 장비", "300",
             "고성능 개발 PC\n테스트 환경 구축",
             "1회성"],
            ["외부 API 비용", "200",
             "카카오 API, 기상청 API\nSSL 인증서, 도메인",
             "월 20만원 × 10개월"],
            ["마케팅/영업", "400",
             "편의점주 대상 홍보\n데모 매장 운영 지원",
             "Beta 확장 시 집중"],
            ["법인 설립/특허", "200",
             "법인 설립 비용\nSW 특허 출원",
             "1회성"],
            ["기타 운영비", "100",
             "사무용품, 교통비 등",
             ""],
            ["합계", "2,000", "", ""],
        ],
        col_widths=[3.5, 2.5, 6, 4]
    )

    # 3-3-2
    add_heading_styled(doc, "3-3-2. 대응자금 사용 계획", level=3)

    add_table(doc,
        ["항목", "금액(만원)", "용도"],
        [
            ["인건비 (대표자)", "500", "시스템 개발/유지보수 (10개월)"],
            ["테스트 매장 운영", "300", "자체 3개 매장 운영비 지원"],
            ["합계", "800", ""],
        ],
        col_widths=[5, 3, 8]
    )

    # 3-3-3
    add_heading_styled(doc, "3-3-3. 매출 계획", level=3)

    add_table(doc,
        ["구분", "Q1 (3~5월)", "Q2 (6~8월)", "Q3 (9~11월)", "Q4 (12월)"],
        [
            ["매장 수", "3개 (자체)", "13개", "30개", "50개"],
            ["월매출", "0원\n(자체 검증)", "39~65만원",
             "90~150만원", "150~250만원"],
            ["누적매출", "0원", "117~195만원",
             "387~645만원", "537~895만원"],
        ],
        col_widths=[2.5, 3, 3.5, 3.5, 3.5]
    )

    # 3-3-4
    add_heading_styled(doc, "3-3-4. BEP 분석", level=3)

    add_para(doc,
        "월 고정비 약 120만원 (서버 80 + API 20 + 기타 20) 기준, "
        "Basic 플랜 3만원/매장 기준 BEP = 40개 매장. "
        "3단계(10~12월) 50개 매장 달성 시 흑자 전환 예상.",
        size=10)

    doc.add_page_break()

    # ════════════════════════════════════════
    # 4. 조직구성
    # ════════════════════════════════════════
    add_heading_styled(doc, "4. 조직구성", level=1)

    # 4-1-1. 대표자 역량
    add_heading_styled(doc, "4-1-1. 대표자 역량", level=2)

    add_table(doc,
        ["항목", "내용"],
        [
            ["성명", "[대표자명]"],
            ["역할", "시스템 설계/개발/운영 총괄 (1인 개발)"],
            ["핵심 역량",
             "CU 편의점 실 운영 경험 (다매장)\n"
             "Python AI/ML 풀스택 개발 역량\n"
             "현장 문제 → 기술 솔루션 전환 능력"],
            ["개발 실적",
             "본 시스템 전체 독자 개발 (2026.01~)\n"
             "5계층 클린 아키텍처 설계 + 구현\n"
             "1,268개 테스트 코드 작성/통과\n"
             "3개 매장 실전 운영 검증 완료"],
            ["기술 스택",
             "Python 3.12, Selenium, SQLite, Flask\n"
             "scikit-learn (RF, GB), 넥사크로 자동화"],
        ],
        col_widths=[3, 13]
    )

    # 4-1-2. 향후 인력 계획
    add_heading_styled(doc, "4-1-2. 향후 인력 계획", level=2)

    add_table(doc,
        ["시기", "인원", "역할", "채용 방식"],
        [
            ["1단계 (3~6월)", "대표자 1인", "개발/운영 전체", "–"],
            ["2단계 (7~9월)", "+1명", "프론트엔드/UX", "인턴 또는 파트타임"],
            ["3단계 (10~12월)", "+1명", "영업/고객지원", "정규직"],
        ],
        col_widths=[3.5, 2, 5, 5.5]
    )

    # 4-2. 보유 지적재산권
    add_heading_styled(doc, "4-2. 보유 지적재산권 및 기술적 성과", level=2)

    add_table(doc,
        ["구분", "내용", "비고"],
        [
            ["소프트웨어",
             "AI 기반 편의점 자동 발주 시스템\n(BGF Auto Order System)",
             "저작권 등록 예정"],
            ["핵심 알고리즘",
             "9단계 예측 파이프라인\n15개 카테고리 독립 Strategy",
             "특허 출원 검토"],
            ["데이터셋",
             "3개 매장 실전 운영 데이터\n(판매/재고/발주/폐기 이력)",
             "2026.01~ 축적 중"],
            ["기술 문서",
             "시스템 아키텍처 문서\n스킬 가이드 6종",
             "CLAUDE.md 등"],
            ["테스트 자산",
             "1,268개 자동화 테스트 스위트",
             "회귀 방지 + 품질 보증"],
        ],
        col_widths=[3, 7, 6]
    )

    return doc


def main():
    doc = create_document()

    # ── 저장 경로 결정 ──
    primary = "/mnt/user-data/outputs/사업계획서_창업중심대학_지역기반_v4.docx"
    fallback = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "..", "..",
        "사업계획서_창업중심대학_지역기반_v4.docx"
    )
    fallback = os.path.normpath(fallback)

    saved = False
    for path in [primary, fallback]:
        try:
            os.makedirs(os.path.dirname(path), exist_ok=True)
            doc.save(path)
            print(f"[OK] 저장 완료: {path}")
            saved = True
            break
        except Exception as e:
            print(f"[WARN] {path} 저장 실패: {e}")

    if not saved:
        # 바탕화면 폴백
        desktop = os.path.expanduser("~/Desktop/사업계획서_창업중심대학_지역기반_v4.docx")
        try:
            doc.save(desktop)
            print(f"[OK] 바탕화면 저장: {desktop}")
        except Exception as e2:
            print(f"[ERROR] 저장 실패: {e2}")
            sys.exit(1)

    return doc


if __name__ == "__main__":
    main()
